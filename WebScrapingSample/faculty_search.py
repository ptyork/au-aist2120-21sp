# REQUIREMENTS (replace pip with pip3 on Mac/Linux/Mimir):
#   pip install requests
#   pip install bs4
#   pip install python-docx

from pprint import pformat # pprint to a string instead of the screen

import logging
import argparse
import requests
import bs4
import json
import docx

# ******************************************************************************
# ARGPARSE MODULE
# ******************************************************************************

# argparse is a super-nice way of "interpreting" the arguments passed in
# to the Python script. Use to define a Command Line Interface (CLI) for your 
# script. Good article here:
#   https://realpython.com/command-line-interfaces-python-argparse/
# Run:
#   python faculty_search.py --help
# to see what this does for you.

parser = argparse.ArgumentParser(
    prog = "faculty_search.py",
    description =
'''
A simple script to pull one or more faculty profiles. By default, this will
output faculty details to the screen as text. However, you can specify 
arguments to output to JSON OR as a Microsoft Word document...potentially 
including the faculty profile image. If more than one profile is found, then
all matching faculty will be processed.
'''
)

parser.add_argument(
    "query",
    type = str,
    help = "The search string used to find the faculty member. If two values are passed, you must use quotation marks and we assume you are asking for a first and last name. If only one value, it is assumed to be the last name."
)

# This adds a "boolean" flag argument (exists or doesn't)
parser.add_argument(
    "--quiet",
    "-q",
    help = "Do not output anything to the screen (only store files as specified in other arguments).",
    action = "store_true"
)

# This adds a "boolean" flag argument (exists or doesn't)
parser.add_argument(
    "--json",
    "-j",
    metavar = "filename",
    help = "Output as JSON to the specified filename",
    type = str
)

# This adds an OPTIONAL named argument string...must set the flag AND pass an string
# In this case -w or --word PLUS the file name
parser.add_argument(
    "--word",
    "-w",
    metavar = "filename",
    help = "Output as an MS Word document to the specified filename",
    type = str
)

# This adds a "boolean" flag argument (exists or doesn't)
parser.add_argument(
    "--image",
    "-i",
    help = "Include the faculty profile image (only if creating word document)",
    action = "store_true"
)

# This adds a "boolean" flag argument (exists or doesn't)
parser.add_argument(
    "--debug",
    "-d",
    help = "Show debug and informational output.",
    action = "store_true"
)

args = parser.parse_args()  # this is the "magic" function that interprets all 
                            # of the arguments, including --help.

if args.debug:
    logging.basicConfig(level=logging.DEBUG)

# debug print the arguments
logging.debug('QUERY:\t' + str(args.query))
logging.debug('DEBUG:\t' + str(args.debug))
logging.debug('QUIET:\t' + str(args.quiet))
logging.debug('JSON:\t' + str(args.json))
logging.debug('WORD:\t' + str(args.word))
logging.debug('IMAGE:\t' + str(args.image))

# insure that they've at least passed in a query
if len(args.query) < 1:
    parser.error("You must supply a search string")

# Split the search string to find last and (maybe) first name
parts = args.query.split()
if len(parts) == 1:
    firstq = None
    lastq = parts[0]
else:
    firstq = parts[0]
    lastq = parts[1]

logging.debug('FIRSTQ:\t' + str(firstq))
logging.debug('LASTQ:\t' + str(lastq))

# ******************************************************************************
# SCRAPING WITH REQUESTS AND BEAUTIFULSOUP (BS4) MODULES
# ******************************************************************************

# OKAY, now first lets figure out what we're doing. We want to search for a
# faculty profile here:
#   https://www.augusta.edu/faculty/directory/
# 
# But issuing a search takes you here:
#   https://www.augusta.edu/faculty/directory/search.php
#
# In order to FIND someone you have to pass in a "query string". A query string
# simply allows you to pass arguments to a URL (like method parameters). It will
# look similar to the following:
#   https://www.augusta.edu/faculty/directory/search.php?last=york
# OR
#   https://www.augusta.edu/faculty/directory/search.php?first=paul&last=york
# NOTE the ? and & separating name=value pairs for these parameters. We'll need
# to add those "properly" using the requests module.
# 
# A page will be returned HOPEFULLY with one or more sections of HTML code that
# looks like the following:

# <div class="collapse-card">
#   <div class="collapse-card__heading">
#     <div class="collapse-card__title">
#       <em class="fa fa-plus fa-fw">&nbsp;</em>
#       <span class="name">York, Paul</span>
#     </div>
#   </div>
#   <div class="collapse-card__body">
#     <div class="row">
#       <div class="col-md-2 text-center"><img class="img-polaroid img-responsive" src="images/faculty/pyork1.jpg"/ alt="York, Paul"></div>
#       <div class="col-md-6">
#         <p class="admrank">Assistant Professor</p>
#         <span style="display: none;" class="col">School of Computer and Cyber Sciences</span>
#         <span style="display: none;" class="dept"></span>
#         <h3>Academic Appointment(s)</h3>
#         <div class="appts">
#           <div class="profession"><p>School of Computer and Cyber Sciences</p></div>
#         </div>
#       </div>
#       <div class="col-md-4 cfooter">
#         <p></p>
#         <p><i class="fa fa-phone fa-fw text-muted"></i> (706) 667-4538</p>
#         <p class="text-muted"><i class="fa fa-map-marker fa-fw"></i> University Hall Cyber Institute</p>
#         <p><a href="view.php?id=pyork1" class="btn btn-primary" style="color: #fff; text-transform: uppercase;">Full Profile</a></p>
#       </div>
#     </div>
#   </div>
# </div>

# So our task will be to retrieve the HTML. Then we will parse this using 
# BeautifulSoup and find the relevant items.

# First let's construct the proper URL (we'll use this URL_BASE later...trust me)
URL_BASE = "https://www.augusta.edu/faculty/directory/"
URL = URL_BASE + "search.php"

# Requests can take "arguments" (or "parameters") as a dictionary if we call
# the request.get() using an extra argument. So let's create the dictionary.

query = { "last": lastq }
if firstq:
    query["first"] = firstq

# Now let's use the requests module to pull down the HTML from the server

try:
    response = requests.get(URL, params=query)
    response.raise_for_status()
except Exception as ex:
    logging.error("An error occurred while executing web search:", ex)
    exit()

# Annnnddd...here we parse the HTML using BeautifulSoup

soup = bs4.BeautifulSoup(response.text, 'html.parser')

# Once parsed, we can search the soup for what we need. First, find all 
# "collapse-card" div's (i.e., something with a class of "collapse-card")

cards = soup.select('.collapse-card')

# If no cards, then no faculty found. Exit.

if not cards:
    print(f'No faculty found matching {firstq} {lastq}')
    exit()

# If found, then we can iterate over each result. Each "card" represents one
# faculty member. We'll create a new list (here a list of dictionaries) to store
# details about each faculty member.
faculty_list = []

for card in cards:
    fac = {}    # each faculty member will be stored as a dictionary

    # Each of the cards is the above HTML...but it is ALSO another soup object.
    # This is really cool because we can now search inside here to find what
    # we want for each faculty member.
    name = card.select_one('.name')
    if name and name.text.strip():
        fac["name"] = name.text.strip()
    rank = card.select_one('.admrank')
    if rank and rank.text.strip():
        fac["rank"] = rank.text.strip()
    college = card.select_one('.col')
    if college and college.text.strip():
        fac["college"] = college.text.strip()
    department = card.select_one('.dept')
    if department and department.text.strip():
        fac["department"] = department.text.strip()

    # THE IMAGE AND LINK TO PROFILE ARE ATTRIBUTES OF ELEMENTS, NOT THE TEXT.
    # SO HANDLED JUST SLIGHTLY DIFFERENTLY. HERE IS WHERE I PROMISED WE WOULD
    # NEED THAT URL_BASE. IT IS NEEDED BECAUSE THE LINKS TO THE IMAGE AND THE
    # PROFILE ARE "RELATIVE" TO THIS BASE URL.

    img = card.select_one('img')
    if img:
        fac["image"] = URL_BASE + img.attrs['src'].strip()
    profile_link = card.select_one('a')
    if profile_link:
        fac["profile"] = URL_BASE + profile_link.attrs['href'].strip()

    # Phone and email appear as text **AFTER** elements that we can find. This
    # is a bit different, but doable. Here we find an element by a class name
    # and then find it's next sibling (i.e. a "brother" or "sister" from the
    # same "parent" element). What is really odd, though, is that the next
    # sibling is just a string, NOT another element.
    # 
    # <p>
    #   <i class="fa fa-phone fa-fw text-muted"></i>
    #   (706) 667-4538
    # </p>
    # 
    # See that the phone number is the next item after the <i></i> element?
    # So instead of getting it's "text" property, it IS TEXT already. We just
    # need to strip away any leading or trailing whitespace characters.
    # 
    # Office location is similar.

    phone_icon = card.select_one('.fa-phone')
    if phone_icon:
        fac["phone"] = phone_icon.next_sibling.strip()
    loc_icon = card.select_one('.fa-map-marker')
    if loc_icon:
        fac["location"] = loc_icon.next_sibling.strip()


    # APPOINTMENTS ARE A BIT TRICKY, SO HANDLED SEPARATELY
    fac["appointments"] = []    # faculty can have many appointments, so this
                                # dictionary key must have a list of strings as
                                # its value...one string for each appointment
    appointments = card.select('.profession')   # will select (maybe) many
    for appt in appointments:
        # fac["appointments"].append(appt.text)   # for each, get the text and add to list
        # THE ABOVE SHOULD WORK FOR ME, BUT FOR SOME FACULTY, THIS "PROFESSION" 
        # SECTION HAS MULTIPLE PARAGRAPHS (college and department). SOOOOO,
        # WE'LL FURTHER SEARCH FOR EACH PARAGRAPH AND "JOIN" THEM. BUT ONLY
        # IF THE PARAGRAPH CONTAINS ACTUAL TEXT, 'CAUSE FOR SOME REASON SOME
        # ARE EMPTY. YUCK!!!!
        all_appts = []
        paras = appt.select('p')
        for p in paras:
            ptext = p.text.strip()
            if ptext:
                all_appts.append(ptext)
        all_appts_string = ' / '.join(all_appts)
        fac["appointments"].append(all_appts_string)
        # WHEW!!! THAT WAS A PAIN

    # FINALLY, add the finished fac dictionary to the faculty_list
    faculty_list.append(fac)

logging.debug(pformat(faculty_list))

# OKAY, SO NOW WE HAVE THE RESULTS STORED IN MEMORY. WE JUST NEED TO OUTPUT THEM
# AS SPECIFIED BY THE COMMAND LINE ARGUMENTS.

# ******************************************************************************
# OUTPUT LIST/DICTIONARY DATA TO SCREEN
# ******************************************************************************

if not args.quiet:
    # PRINT AS PLAIN TEXT
    for fac in faculty_list:
        print("="*80)
        print(f'{"NAME":15}{fac["name"]}')
        if "rank" in fac:
            print(f'{"RANK":15}{fac["rank"]}')
        if "college" in fac:
            print(f'{"COLLEGE":15}{fac["college"]}')
        if "department" in fac:
            print(f'{"DEPARTMENT":15}{fac["department"]}')
        if "appointments" in fac and len(fac["appointments"]) > 0:
            print("APPOINTMENTS")
            for appt in fac["appointments"]: # remember, this is a list of strings
                print(f'  * {appt}')
        if "phone" in fac:
            print(f'{"PHONE":15}{fac["phone"]}')
        if "location" in fac:
            print(f'{"LOCATION":15}{fac["location"]}')
        if "profile" in fac:
            print(f'{"PROFILE":15}{fac["profile"]}')
        print("="*80)
        print()

# ******************************************************************************
# OUTPUT LIST/DICTIONARY DATA TO JSON
# ******************************************************************************

if args.json:
    # WOW, ISN'T JSON SUPER-EASY?!?!?!?!
    logging.info(f"SAVING JSON TO {args.json}")
    try:
        with open(args.json, 'w') as fi:
            json.dump(faculty_list, fi, indent=4) # indent here is just to make it pretty
    except Exception as ex:
        logging.error(f"An unexpected error occurred while saving {args.json}: {ex}")

# ******************************************************************************
# OUTPUT LIST/DICTIONARY DATA TO MICROSOFT WORD
# ******************************************************************************

if args.word:
    # WORD IS MAYBE __NOT__ AS EASY AS JSON...
    logging.info(f"SAVING JSON TO {args.json}")

    # Righteously swiped this function from here:
    # https://stackoverflow.com/questions/41578438/how-do-you-keep-table-rows-together-in-python-docx
    # Needed because tables created below were inserting page breaks in really
    # awkward spots.
    # NOTE: Do as I say, not as I do...normally DO NOT define functions inside
    #       of an if-block like this. I'm JUST doing it because it is easier
    #       for you do understand and "borrow" things if they are in
    #       self-contained blocks like this.
    def keep_table_on_one_page(doc):
        tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
        for tag in tags:
            ppr = tag.get_or_add_pPr()
            ppr.keepNext_val = True

    # Here is the real code to save the faculty_list as a "pretty" Word document
    try:
        # Create a new document
        doc = docx.Document()

        # Store some standard "styles" in variables for convenience
        tableStyle = doc.styles['Medium Grid 2 Accent 1']
        nameStyle = doc.styles['Heading 1']
        listStyle = doc.styles['List Bullet']

        # Main heading
        doc.add_heading('Faculty Search Results', 0)
        
        # Create and populate a Summary table
        doc.add_heading('Search Summary', level=1)
        table = doc.add_table(3,2)
        table.style = tableStyle
        table.autofit = False
        table.cell(0,0).text = 'First Name'
        table.cell(0,0).width = 1100000
        table.cell(0,1).text = str(firstq)
        table.cell(0,1).width = 44500000
        table.cell(1,0).text = 'Last Name'
        table.cell(1,0).width = 1100000
        table.cell(1,1).text = str(lastq)
        table.cell(1,1).width = 4450000
        table.cell(2,0).text = 'Result Count'
        table.cell(2,0).width = 1100000
        table.cell(2,1).text = str(len(faculty_list))
        table.cell(2,1).width = 4450000

        # NOTE: I couldn't figure out an easy way to apply a width to a specific
        #       column and instead had to apply it to each individual cell.
        #       That's a pain...as are a few things with this docx module.

        # Now print each faculty member
        doc.add_heading('Results', level=1)

        for fac in faculty_list:
            # Each will be in a table
            table = doc.add_table(1,3)
            table.style = tableStyle
            table.autofit = False
            # The header row is a heading 1, BUT it's in a single cell of this
            # table "merged" from three separate cells on the first row
            table.cell(0,0).text = str(fac["name"])
            table.cell(0,0).merge(table.cell(0,2))
            table.cell(0,0).paragraphs[0].style = nameStyle

            # Now just add rows to the table for each attribute
            if "rank" in fac:
                row = table.add_row()
                row.cells[0].text = "Rank"
                row.cells[1].text = str(fac["rank"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "college" in fac:
                row = table.add_row()
                row.cells[0].text = "College"
                row.cells[1].text = str(fac["college"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "department" in fac:
                row = table.add_row()
                row.cells[0].text = "Department"
                row.cells[1].text = str(fac["department"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "appointments" in fac and len(fac["appointments"]) > 0:
                row = table.add_row()
                row.cells[0].text = "Appointment(s)"
                # This one's a bit different since the appointments dictionary
                # value is a LIST of strings. So, we'll add each as separate 
                # "bullets" (or rather bullet-style paragraphs)
                row.cells[1].text = fac["appointments"][0]
                for appt in fac["appointments"][1:]:
                    row.cells[1].add_paragraph(appt)
                for p in row.cells[1].paragraphs:
                    p.style = listStyle
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "phone" in fac:
                row = table.add_row()
                row.cells[0].text = "Phone Number"
                row.cells[1].text = str(fac["phone"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "location" in fac:
                row = table.add_row()
                row.cells[0].text = "Location"
                row.cells[1].text = str(fac["location"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            if "profile" in fac:
                row = table.add_row()
                row.cells[0].text = "Profile Link"
                row.cells[1].text = str(fac["profile"])
                row.cells[0].width = 1100000
                row.cells[1].width = 3300000
            
            # Optionally retrieve and insert the image
            if args.image and "image" in fac:
                # This import allows me to "stream" the image directly into
                # Word. Normally I'd have this at the top, but it's a bit 
                # confusing, so I put it here to minimize this.
                # https://2.python-requests.org/en/master/user/quickstart/#binary-response-content
                from io import BytesIO

                # merge all cells in column 3
                col3 = table.column_cells(2)
                col3[1].merge(col3[-1])
                col3[1].width = 1150000

                try:
                    response = requests.get(fac["image"])
                    response.raise_for_status()
                    # This is weird. I have to add a "run" to the paragraph in
                    # the table cell in order to add a picture to it.
                    # https://github.com/python-openxml/python-docx/issues/216
                    run = col3[1].paragraphs[0].add_run()
                    # And finally "stream" the picture into the document
                    run.add_picture(BytesIO(response.content), width=1000000)
                except requests.HTTPError as ex:
                    logging.error(f'Unable to retrieve requested resource: {ex}')
                except Exception as ex:
                    logging.error(f'An unexpected error has occurred: {ex}')
            else:
                # No image so "merge" columns 2 & 3
                for row in table.rows[1:]:
                    row.cells[1].merge(row.cells[2])
                    row.cells[1].width = 4450000
                
            doc.add_paragraph() # space between result tables
        
        keep_table_on_one_page(doc)
        doc.save(args.word)

    except Exception as ex:
        logging.error(f"An unexpected error occurred while saving {args.word}: {ex}")


